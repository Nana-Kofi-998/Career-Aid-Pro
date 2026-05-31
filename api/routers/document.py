"""Document generation routes."""

from __future__ import annotations

from io import BytesIO
import html
import re

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse

from api.deps import get_current_user
from api.schemas import DocumentGenerateRequest, DocumentFromChatRequest

router = APIRouter(prefix="/document", tags=["document"])


def _clean_inline_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def _split_title_and_body(text: str) -> tuple[str, list[str]]:
    lines = [line.rstrip() for line in text.strip().splitlines()]
    while lines and not lines[0].strip():
        lines.pop(0)
    if not lines:
        return "Career Aid Document", []

    first = _clean_inline_markdown(lines[0]).strip("#: ")
    if len(first) <= 120 and (
        first.lower().startswith(("career aid", "becoming", "my ", "report", "plan"))
        or lines[0].startswith("#")
    ):
        return first or "Career Aid Document", lines[1:]
    return "Career Aid Document", lines


def _iter_blocks(lines: list[str]):
    paragraph: list[str] = []

    def flush_paragraph():
        nonlocal paragraph
        if paragraph:
            yield ("paragraph", _clean_inline_markdown(" ".join(paragraph)))
            paragraph = []

    for raw in lines:
        line = raw.strip()
        if not line:
            yield from flush_paragraph()
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        label_heading = re.match(r"^\*\*(.+?):\*\*\s*(.*)$", line)
        numbered_heading = re.match(r"^(\d+)\.\s+\*\*(.+?)\*\*:?\s*(.*)$", line)
        bullet = re.match(r"^[*-]\s+(.+)$", line)

        if heading:
            yield from flush_paragraph()
            yield ("heading", _clean_inline_markdown(heading.group(2)))
        elif numbered_heading:
            yield from flush_paragraph()
            title = f"{numbered_heading.group(1)}. {_clean_inline_markdown(numbered_heading.group(2))}"
            yield ("heading", title)
            if numbered_heading.group(3).strip():
                yield ("paragraph", _clean_inline_markdown(numbered_heading.group(3)))
        elif label_heading:
            yield from flush_paragraph()
            yield ("heading", _clean_inline_markdown(label_heading.group(1)))
            if label_heading.group(2).strip():
                yield ("paragraph", _clean_inline_markdown(label_heading.group(2)))
        elif bullet:
            yield from flush_paragraph()
            yield ("bullet", _clean_inline_markdown(bullet.group(1)))
        else:
            paragraph.append(line)

    yield from flush_paragraph()


def generate_word_doc(text: str) -> bytes:
    """Generate a Word document from plain text."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        title_text, body_lines = _split_title_and_body(text)
        title = doc.add_heading(title_text, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for block_type, content in _iter_blocks(body_lines):
            if not content:
                continue
            if block_type == "heading":
                doc.add_heading(content, level=2)
            elif block_type == "bullet":
                doc.add_paragraph(content, style="List Bullet")
            else:
                paragraph = doc.add_paragraph(content)
                paragraph.style = "Normal"
                for run in paragraph.runs:
                    run.font.size = Pt(11)

        byte_io = BytesIO()
        doc.save(byte_io)
        byte_io.seek(0)
        return byte_io.read()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate Word document: {str(e)}")


def generate_pdf_doc(text: str) -> bytes:
    """Generate a PDF document from plain text."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_CENTER, TA_LEFT

        byte_io = BytesIO()
        doc = SimpleDocTemplate(byte_io, pagesize=letter,
                              rightMargin=inch, leftMargin=inch,
                              topMargin=inch, bottomMargin=inch)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            leading=24,
            spaceAfter=20,
            alignment=TA_CENTER
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            leading=18,
            spaceBefore=12,
            spaceAfter=6,
            textColor="#0f3b46",
        )
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=12,
            leading=16,
            spaceAfter=8,
            alignment=TA_LEFT
        )
        bullet_style = ParagraphStyle(
            'CustomBullet',
            parent=normal_style,
            leftIndent=18,
            firstLineIndent=-10,
        )

        title_text, body_lines = _split_title_and_body(text)
        story = [Paragraph(html.escape(title_text), title_style)]
        story.append(Spacer(1, 12))

        for block_type, content in _iter_blocks(body_lines):
            if not content:
                continue
            if block_type == "heading":
                story.append(Paragraph(html.escape(content), heading_style))
            elif block_type == "bullet":
                story.append(Paragraph(f"&bull; {html.escape(content)}", bullet_style))
            else:
                story.append(Paragraph(html.escape(content), normal_style))

        doc.build(story)
        byte_io.seek(0)
        return byte_io.read()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF document: {str(e)}")


@router.post("/generate-word")
async def generate_word(
    payload: DocumentGenerateRequest,
    user: dict = Depends(get_current_user)
):
    """Generate and return a Word document."""
    try:
        doc_bytes = generate_word_doc(payload.text)
        return StreamingResponse(
            BytesIO(doc_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=document.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/generate-pdf")
async def generate_pdf(
    payload: DocumentGenerateRequest,
    user: dict = Depends(get_current_user)
):
    """Generate and return a PDF document."""
    try:
        pdf_bytes = generate_pdf_doc(payload.text)
        return StreamingResponse(
            BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=document.pdf"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/generate-from-chat")
async def generate_from_chat(
    payload: DocumentFromChatRequest,
    user: dict = Depends(get_current_user)
):
    """Generate a document from a specific chat."""
    from career_aid_pro import database as db
    
    chat_id = payload.chat_id
    doc_type = payload.doc_type
    
    # Verify chat belongs to user
    chat = db.load_chat(chat_id, user["username"])
    if not chat:
        raise HTTPException(status_code=404, detail="Chat not found")
    
    # Format chat history as text
    history = chat.get("history", [])
    text_lines = ["Career-Aid Pro Chat Export"]
    text_lines.append("=" * 50)
    text_lines.append(f"Chat ID: {chat_id}")
    text_lines.append(f"Mode: {chat.get('mode', 'Unknown')}")
    text_lines.append(f"Created: {chat.get('created_at', 'Unknown')}")
    text_lines.append("")
    text_lines.append("Conversation:")
    text_lines.append("-" * 30)
    
    for msg in history:
        role = msg.get("role", "unknown")
        content = msg.get("content", "")
        timestamp = msg.get("ts", "")
        text_lines.append(f"[{timestamp}] {role.title()}: {content}")
        text_lines.append("")
    
    text_content = "\n".join(text_lines)
    
    if doc_type.lower() == "pdf":
        doc_bytes = generate_pdf_doc(text_content)
        media_type = "application/pdf"
        filename = f"chat_{chat_id}.pdf"
    else:
        doc_bytes = generate_word_doc(text_content)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = f"chat_{chat_id}.docx"
    
    return StreamingResponse(
        BytesIO(doc_bytes),
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )
